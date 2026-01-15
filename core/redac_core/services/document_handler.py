"""Document handling service for PDF, DOCX, and TXT files."""

import io
import os
import tempfile
import logging
from pathlib import Path
from typing import BinaryIO, List

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


class DocumentHandler:
    """Handler for different document types."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}
    MAX_FILE_SIZE_MB = 50

    def __init__(self):
        self.temp_dir = tempfile.mkdtemp(prefix="redactx_")
        logger.info(f"DocumentHandler initialized with temp_dir: {self.temp_dir}")

    def is_supported(self, filename: str) -> bool:
        """Check if file extension is supported."""
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS

    def get_file_type(self, filename: str) -> str:
        """Get normalized file type from filename."""
        ext = Path(filename).suffix.lower()
        if ext in {".doc", ".docx"}:
            return "docx"
        return ext.lstrip(".")

    def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text content from a document."""
        file_type = self.get_file_type(filename)

        if file_type == "pdf":
            return self._extract_from_pdf(file_content)
        elif file_type == "docx":
            return self._extract_from_docx(file_content)
        elif file_type in {"txt", "md"}:
            return self._extract_from_text(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF."""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

            doc.close()
            
            extracted = "\n\n".join(text_parts)
            if not extracted.strip():
                raise ValueError("No text found in PDF. Scanned documents (images) are not supported.")
                
            return extracted
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            raise ValueError(f"Failed to process PDF: {e}")

    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            extracted = "\n\n".join(paragraphs)
            
            if not extracted.strip():
                raise ValueError("No text found in DOCX.")
                
            return extracted
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise ValueError(f"Failed to process DOCX: {e}")

    def _extract_from_text(self, content: bytes) -> str:
        """Extract text from plain text file."""
        try:
            # Try UTF-8 first, then latin-1 as fallback
            text = ""
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                text = content.decode("latin-1")
                
            if not text.strip():
                raise ValueError("File is empty.")
                
            return text
        except Exception as e:
            logger.error(f"Failed to decode text file: {e}")
            raise ValueError(f"Failed to process text file: {e}")

    def create_redacted_document(
        self,
        original_content: bytes,
        filename: str,
        redacted_text: str,
        output_format: str | None = None,
        entities: List = None,
        method: str = "mask"
    ) -> tuple[bytes, str]:
        """Create a redacted version of the document.

        Args:
            original_content: Original file content
            filename: Original filename
            redacted_text: Redacted text to use
            output_format: Force output format (txt, json) or None for same as input
            entities: Optional list of detected entities for in-place redaction
            method: Redaction method (mask, remove, replace, hash)

        Returns:
            Tuple of (redacted_content, new_filename)
        """
        file_type = self.get_file_type(filename)
        stem = Path(filename).stem
        
        # Default entities to empty list if None
        entities = entities or []

        # If output format is specified, override file type
        if output_format and output_format != "same":
            file_type = output_format

        if file_type == "pdf":
            # If we have entities and use in-place, we prioritize that to preserve layout
            # Only use reconstruction if explicitly requested or no entities available to map
            return self._create_redacted_pdf(original_content, redacted_text, stem, entities, method)
        elif file_type == "docx":
            return self._create_redacted_docx(redacted_text, stem)
        elif file_type in {"txt", "md"}:
            ext = ".md" if file_type == "md" else ".txt"
            return redacted_text.encode("utf-8"), f"{stem}_redacted{ext}"
        elif file_type == "json":
            return redacted_text.encode("utf-8"), f"{stem}_redacted.json"
        else:
            # Default to txt
            return redacted_text.encode("utf-8"), f"{stem}_redacted.txt"

    def _create_redacted_pdf(
        self, 
        original_content: bytes, 
        redacted_text: str, 
        stem: str, 
        entities: List = None,
        method: str = "mask"
    ) -> tuple[bytes, str]:
        """Create a new PDF with redacted text.
        
        If entities are provided, attempts in-place redaction to preserve layout.
        Otherwise falls back to text-reconstruction.
        """
        if not entities:
            # If no PII detected (e.g. high confidence threshold), return original document
            # to preserve layout perfectly rather than reconstructing plain text.
            return original_content, f"{stem}_redacted.pdf"

        try:
            doc = fitz.open(stream=original_content, filetype="pdf")
            
            # Strategy 1: In-Place Redaction (Preserves Layout)
            # We already checked 'if entities' above, so we proceed.
            if entities:
                # Group entities by text
                unique_texts = set(e.text for e in entities if e.text and e.text.strip())
                
                # Determine colors based on method
                fill_color = (0, 0, 0) # Default Black
                text_color = None
                
                if method == "remove":
                    fill_color = (1, 1, 1) # White (matches paper)
                elif method in ("replace", "hash"):
                    fill_color = (0.9, 0.9, 0.9) # Light Gray
                    text_color = (0.3, 0.3, 0.3) # Dark Gray Text

                for page in doc:
                    for entity in entities:
                        if not entity.text or not entity.text.strip():
                            continue
                            
                        # Search for this specific entity text
                        hits = page.search_for(entity.text)
                        for rect in hits:
                            # Configure annotation based on method
                            annot_text = None
                            if method == "replace":
                                annot_text = f"[{entity.label}]"
                            elif method == "hash":
                                import hashlib
                                annot_text = hashlib.md5(entity.text.encode()).hexdigest()[:8]

                            page.add_redact_annot(
                                rect, 
                                fill=fill_color,
                                text=annot_text,
                                text_color=text_color,
                                fontsize=8
                            )
                    
                    page.apply_redactions()
                
                output = doc.tobytes()
                doc.close()
                return output, f"{stem}_redacted.pdf"
            
            # Strategy 2: Reconstruction (Text Only) - Fallback
            doc.close()
            
            doc = fitz.open()
            
            # Constants
            fontsize = 10
            fontname = "helv"
            margin = 50
            line_height = 12
            
            # Simple text insertion logic handling simple pagination
            # We split by lines to manage basic flow
            lines = redacted_text.split('\n')
            
            current_page = doc.new_page()
            y = margin
            page_height = current_page.rect.height
            
            shape = current_page.new_shape()
            
            for line in lines:
                if y > page_height - margin:
                    # New Page
                    shape.commit()
                    current_page = doc.new_page()
                    shape = current_page.new_shape()
                    y = margin
                
                # Check if line contains page break marker
                if line.startswith("--- Page") and line.endswith("---"):
                    pass

                # Insert text
                current_page.insert_text((margin, y), line, fontsize=fontsize, fontname=fontname)
                y += line_height
            
            # Save to bytes
            output = doc.tobytes()
            doc.close()

            return output, f"{stem}_redacted.pdf"
        except Exception as e:
            logger.error(f"Failed to create redacted PDF: {e}")
            # Fallback to text file
            return redacted_text.encode("utf-8"), f"{stem}_redacted.txt"

    def _create_redacted_docx(
        self, redacted_text: str, stem: str
    ) -> tuple[bytes, str]:
        """Create a new DOCX with redacted text."""
        try:
            doc = Document()

            # Split text into paragraphs and add to document
            paragraphs = redacted_text.split("\n\n")
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)

            # Save to bytes
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)

            return output_buffer.read(), f"{stem}_redacted.docx"
        except Exception as e:
            logger.error(f"Failed to create redacted DOCX: {e}")
            # Fallback to text file
            return redacted_text.encode("utf-8"), f"{stem}_redacted.txt"

    def cleanup(self):
        """Clean up temporary files."""
        import shutil

        try:
            shutil.rmtree(self.temp_dir, ignore_errors=True)
            logger.info(f"Cleaned up temp directory: {self.temp_dir}")
        except Exception as e:
            logger.error(f"Failed to cleanup temp directory: {e}")


# Singleton instance
_document_handler: DocumentHandler | None = None


def get_document_handler() -> DocumentHandler:
    """Get document handler instance."""
    global _document_handler
    if _document_handler is None:
        _document_handler = DocumentHandler()
    return _document_handler
